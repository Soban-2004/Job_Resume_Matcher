from io import BytesIO

import fitz  # PyMuPDF
from docx import Document


def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    # Resumes commonly use tables for column layouts (skills grids, date/role
    # alignment) -- docx2txt included table text by default, so this keeps
    # feature parity with python-docx's paragraph-only default.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def load_document(filename: str, content: bytes) -> str | None:
    """Extracts text from an uploaded PDF/DOCX/TXT file, entirely in-memory.

    (The old Streamlit version wrote .docx uploads to a shared temp.docx
    path on disk, which breaks under concurrent requests.)
    """
    try:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            with fitz.open(stream=content, filetype="pdf") as doc:
                return "".join(page.get_text() for page in doc)
        elif lower.endswith(".docx"):
            return _extract_docx_text(content)
        elif lower.endswith(".txt"):
            return content.decode("utf-8", errors="ignore")
        return None
    except Exception:
        return None
